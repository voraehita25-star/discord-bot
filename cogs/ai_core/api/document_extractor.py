"""Extract plain text from dashboard document attachments.

First-upload flow: when a user attaches a PDF / DOCX / text file to the
dashboard chat, this module pulls out the text content so it can be
persisted to SQLite and auto-injected into every subsequent AI turn —
even across conversations, sessions, and bot restarts. The original
binary is NOT kept (it's deleted by the caller's temp-dir cleanup),
so storage scales with extracted text length, not PDF page count.

Design notes:
  - PDF extraction uses ``pypdf`` (pure-Python, fast, handles most layouts).
    Tables and complex multi-column layouts may render with wonky spacing,
    but prose and bullet lists come through cleanly — which is what users
    attach for RP campaigns.
  - DOCX uses ``python-docx`` and just concatenates paragraph text.
  - Everything else (``.txt``, ``.md``, ``.json``, code, etc.) is already
    text — we just normalise newlines and return it.
  - No network calls. Extraction runs in the request thread; for big PDFs
    callers may want to wrap with ``asyncio.to_thread`` to avoid blocking
    the event loop.

Extracted text is sanitised to strip control chars + enforce a hard cap
per document so one oversized file can't balloon the DB.
"""

from __future__ import annotations

import base64
import binascii
import logging
import re
import time
from dataclasses import dataclass
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


# python-docx parses DOCX (a ZIP of XML) with lxml. Its parser is built with
# ``resolve_entities=False`` (verified: ``docx.oxml.parser.oxml_parser``) and
# lxml defaults to ``no_network=True``, so external-entity (XXE) and
# billion-laughs expansion are blocked at parse time — pinned by
# ``test_docx_parser_does_not_expand_xml_entities``. Combined with the streaming
# zip-bomb guard in ``_extract_docx`` below, DOCX is safe whenever python-docx
# imports. (A ``defusedxml.lxml`` monkey-patch used to be applied here, but it
# ran *after* ``oxml_parser`` was already built — a no-op — and only emitted a
# DeprecationWarning, since ``defusedxml.lxml`` is deprecated upstream.)
try:
    import docx as _docx_probe  # noqa: F401 - module-level availability probe

    DOCX_DISABLED = False
except ImportError:
    DOCX_DISABLED = True
    logger.warning(
        "python-docx not installed — DOCX extraction is DISABLED. "
        "Install python-docx to enable DOCX support."
    )


# Cap extracted text per file. 500K chars ≈ ~120-200K tokens depending on
# language. This is the max that's worth persisting; anything larger would
# dominate the prompt budget when re-injected.
MAX_EXTRACTED_CHARS = 500_000

# Truncation marker appended when extracted text is sliced. The slice budget
# subtracts this so the final ``len(text) <= MAX_EXTRACTED_CHARS`` invariant
# actually holds — previously the marker pushed the final length ~18 chars
# over the cap, which is harmless in practice but breaks tests that assert
# strict equality on the bound.
_TRUNCATION_MARKER = "\n\n[... truncated]"
_TRUNCATION_BUDGET = MAX_EXTRACTED_CHARS - len(_TRUNCATION_MARKER)


@dataclass(slots=True)
class ExtractedDocument:
    """Result of a successful extraction."""

    filename: str
    kind: str  # 'pdf' | 'docx' | 'text'
    text: str
    char_count: int
    page_count: int | None = None  # only set for paged formats (PDF)


def extract_from_payload(payload: dict[Any, Any]) -> ExtractedDocument | None:
    """Extract text from a single frontend document payload.

    Expected shape (mirrors ``DocumentAttachManager.get()``)::
        {"name": str, "mime": str, "kind": "binary"|"text", "data": str, "size_bytes": int}

    Returns ``None`` if extraction fails (corrupt file, unsupported format,
    empty result). The caller should log / toast and move on — we never
    raise from here because one bad file shouldn't break the whole turn.
    """
    if not isinstance(payload, dict):
        return None
    name = str(payload.get("name", "attachment"))[:200]
    kind_hint = payload.get("kind")
    data_field = payload.get("data")
    if not isinstance(data_field, str) or not data_field:
        return None

    ext = _extension_of(name)

    try:
        if ext == ".pdf":
            return _extract_pdf(name, data_field)
        if ext == ".docx":
            return _extract_docx(name, data_field)
        # Everything else: text. kind_hint from frontend is authoritative
        # for text-vs-binary classification at this point (extension not
        # matching PDF/DOCX + data being string implies text).
        if kind_hint == "text" or not data_field.startswith("data:"):
            return _extract_text(name, data_field)
    except Exception as e:
        logger.warning("Document extraction failed for %s: %s", name, e)
    return None


def _extension_of(filename: str) -> str:
    """Lowercase extension including the leading dot, or empty string."""
    m = re.search(r"\.[A-Za-z0-9]+$", filename)
    return m.group(0).lower() if m else ""


def _decode_data_url(data_field: str) -> bytes | None:
    """Decode a ``data:<mime>;base64,<b64>`` URL to bytes. Returns None
    if the format is wrong or the base64 is corrupt."""
    if "," not in data_field or not data_field.startswith("data:"):
        return None
    _header, _, payload = data_field.partition(",")
    try:
        return base64.b64decode(payload, validate=True)
    except (ValueError, binascii.Error):
        return None


def _extract_pdf(filename: str, data_field: str) -> ExtractedDocument | None:
    """Decode base64 PDF and pull text via pypdf.

    Failures (malformed base64, encrypted PDF, corrupt stream) return None
    so callers can fall back to "only saw it this turn" semantics instead
    of erroring out the whole request.
    """
    pdf_bytes = _decode_data_url(data_field)
    if pdf_bytes is None:
        return None

    # Lazy import — pypdf is ~1 MB and we only need it here. Keeps the
    # bot's cold-start cheap for users who never attach PDFs.
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed — PDF text extraction unavailable")
        return None

    try:
        reader = PdfReader(BytesIO(pdf_bytes))
    except Exception as e:
        logger.warning("pypdf failed to parse %s: %s", filename, e)
        return None

    # Skip encrypted PDFs — decryption would need a password we don't have
    # and silently returning empty text would mislead the user.
    if getattr(reader, "is_encrypted", False):
        logger.info("Skipping encrypted PDF: %s", filename)
        return None

    # Reject excessively large PDFs before doing per-page extraction. A
    # malicious or malformed PDF with tens of thousands of pages can spin
    # pypdf for minutes per file; pages we'd never use anyway are bounded
    # by MAX_DOC_PAGES below, but we want to fail fast rather than chew
    # through 50k empty pages first.
    MAX_DOC_PAGES = 2000
    try:
        page_count = len(reader.pages)
    except Exception:
        page_count = 0
    if page_count > MAX_DOC_PAGES:
        logger.warning(
            "Rejecting oversized PDF %s (%d pages > %d cap)",
            filename,
            page_count,
            MAX_DOC_PAGES,
        )
        return None

    pages: list[str] = []
    start = time.monotonic()
    for idx, page in enumerate(reader.pages):
        # Wall-clock guard: a single PDF with thousands of dense pages
        # (or pathological text operators) can keep pypdf busy for minutes
        # and starve the event loop. 60s gives even a 2k-page document a
        # fair shot while still bounding worst-case latency.
        if time.monotonic() - start > 60:
            logger.warning(
                "PDF extraction timeout for %s after page %d, returning partial result",
                filename,
                idx,
            )
            break
        try:
            # "layout" mode reconstructs text using the PDF's positional
            # information, which preserves real line breaks from the source
            # document much better than the default "plain" mode. That
            # matters because pypdf's plain mode emits newlines for every
            # positional jump — indistinguishable from actual line breaks —
            # and any downstream rejoin can't tell real breaks from noise.
            # Layout mode sidesteps the ambiguity.
            text = page.extract_text(extraction_mode="layout") or ""
        except Exception as e:
            # Individual page failures are common (graphic-heavy pages);
            # don't drop the whole doc — just note and move on. Also fall
            # back to plain mode in case layout mode specifically barfs on
            # this page's text operators.
            logger.debug("pypdf layout-mode skipped page %d of %s: %s", idx + 1, filename, e)
            try:
                text = page.extract_text() or ""
            except Exception as e2:
                logger.debug(
                    "pypdf plain-mode also failed on page %d of %s: %s", idx + 1, filename, e2
                )
                continue
        # Order matters: rejoin BEFORE _normalise so paragraph collapse
        # runs on the fixed-up text. _normalise's \n{3,} → \n\n pass
        # handles any extras the rejoin left behind.
        text = _rejoin_pdf_lines(text)
        text = _normalise(text)
        if text:
            pages.append(f"[Page {idx + 1}]\n{text}")

    if not pages:
        return None

    combined = "\n\n".join(pages)
    if len(combined) > MAX_EXTRACTED_CHARS:
        combined = combined[:_TRUNCATION_BUDGET] + _TRUNCATION_MARKER
    return ExtractedDocument(
        filename=filename,
        kind="pdf",
        text=combined,
        char_count=len(combined),
        page_count=len(reader.pages),
    )


def _extract_docx(filename: str, data_field: str) -> ExtractedDocument | None:
    """Decode base64 DOCX and pull paragraph text via python-docx."""
    if DOCX_DISABLED:
        logger.warning("DOCX extraction disabled (python-docx not installed): %s", filename)
        return None

    docx_bytes = _decode_data_url(data_field)
    if docx_bytes is None:
        return None

    try:
        import docx
    except ImportError:
        logger.warning("python-docx not installed — DOCX extraction unavailable")
        return None

    # Zip-bomb guard. DOCX is a ZIP container; python-docx itself does not
    # cap per-entry decompression. A 1 KiB DOCX whose [Content_Types].xml
    # decompresses to multi-GB would blow out the process. The header-
    # declared file_size is forgeable — a malicious DOCX can advertise a
    # tiny size while actually decompressing to gigabytes — so we stream-
    # decompress every entry in 64 KiB chunks and abort the moment the
    # observed bytes exceed the cap. Cap is a hard 50 MiB aggregate.
    import zipfile

    _MAX_DOCX_TOTAL_BYTES = 50 * 1024 * 1024  # 50 MiB aggregate (real bytes)
    _MAX_DOCX_ENTRIES = 1000  # Real DOCX files have ~10-50 entries; 1000 is generous
    _CHUNK = 64 * 1024
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes)) as zf:
            entries = zf.infolist()
            # Cap entry count BEFORE iterating. A crafted DOCX with millions
            # of zero-length entries makes ``infolist()`` allocate
            # proportionally and the per-entry loop spin forever — the
            # byte cap above doesn't trigger because each entry is empty.
            if len(entries) > _MAX_DOCX_ENTRIES:
                logger.warning(
                    "DOCX %s rejected: %d entries exceeds %d-entry cap (zip-bomb guard)",
                    filename,
                    len(entries),
                    _MAX_DOCX_ENTRIES,
                )
                return None
            total = 0
            for info in entries:
                # Defense-in-depth: reject path-traversal entry names.
                # We only read into memory today, but a future change that
                # writes any entry to disk would inherit a traversal sink
                # if we didn't filter at the source.
                entry_name = info.filename or ""
                if (
                    entry_name.startswith(("/", "\\"))
                    or ".." in entry_name.replace("\\", "/").split("/")
                ):
                    logger.warning(
                        "DOCX %s rejected: suspicious entry name %r",
                        filename,
                        entry_name,
                    )
                    return None
                with zf.open(info.filename) as entry:
                    while True:
                        chunk = entry.read(_CHUNK)
                        if not chunk:
                            break
                        total += len(chunk)
                        if total > _MAX_DOCX_TOTAL_BYTES:
                            logger.warning(
                                "DOCX %s exceeded %d-byte decompression cap (zip-bomb guard)",
                                filename,
                                _MAX_DOCX_TOTAL_BYTES,
                            )
                            return None
    except zipfile.BadZipFile:
        logger.warning("DOCX %s is not a valid zip", filename)
        return None
    except Exception:
        logger.exception("DOCX %s zip preflight failed", filename)
        return None

    try:
        document = docx.Document(BytesIO(docx_bytes))
    except Exception as e:
        # python-docx parses with resolve_entities=False, so external-entity /
        # billion-laughs payloads are never expanded (no XXE). A malformed or
        # otherwise unparseable doc lands here and is safely rejected.
        logger.warning("python-docx failed to parse %s: %s", filename, e)
        return None

    # Grab paragraphs + table cells. Tables are flattened cell-by-cell
    # with tab separators — not beautiful, but preserves order and content.
    chunks: list[str] = []
    for para in document.paragraphs:
        text = _normalise(para.text)
        if text:
            chunks.append(text)
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join(_normalise(cell.text) for cell in row.cells)
            if row_text.strip():
                chunks.append(row_text)

    if not chunks:
        return None

    combined = "\n\n".join(chunks)
    if len(combined) > MAX_EXTRACTED_CHARS:
        combined = combined[:_TRUNCATION_BUDGET] + _TRUNCATION_MARKER
    return ExtractedDocument(
        filename=filename,
        kind="docx",
        text=combined,
        char_count=len(combined),
    )


def _extract_text(filename: str, data_field: str) -> ExtractedDocument | None:
    """Text files arrive already decoded as UTF-8 strings — just
    normalise + truncate."""
    text = _normalise(data_field)
    if not text:
        return None
    if len(text) > MAX_EXTRACTED_CHARS:
        text = text[:_TRUNCATION_BUDGET] + _TRUNCATION_MARKER
    return ExtractedDocument(
        filename=filename,
        kind="text",
        text=text,
        char_count=len(text),
    )


def _normalise(text: str) -> str:
    """Strip control chars (keep \\n \\t), collapse runs of whitespace."""
    # Remove C0 controls except \t (\x09) and \n (\x0a), plus DEL (\x7f).
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    # Collapse >2 blank lines to exactly 2 — preserves paragraph breaks
    # without letting one quirky PDF blow up the total char count.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# Script ranges that DON'T use spaces between words. When joining lines
# whose neighbouring characters both belong to one of these scripts, we
# drop the separator entirely instead of turning it into a space — otherwise
# Thai / CJK text comes out as "คำ คำ คำ" (space between every word) which
# reads as broken/pidgin in those languages.
_CONTINUOUS_SCRIPT_RANGES = (
    "฀-๿"  # Thai
    "຀-໿"  # Lao
    "ក-៿"  # Khmer
    "぀-ヿ"  # Hiragana + Katakana
    "ㇰ-ㇿ"  # Katakana phonetic extensions
    "一-鿿"  # CJK Unified Ideographs
    "㐀-䶿"  # CJK Extension A
    "豈-﫿"  # CJK Compatibility Ideographs
    "가-힯"  # Hangul Syllables
)
_RE_CONTINUOUS_CHAR = re.compile(rf"[{_CONTINUOUS_SCRIPT_RANGES}]")
_RE_SPACE_BETWEEN_CONTINUOUS = re.compile(
    rf"([{_CONTINUOUS_SCRIPT_RANGES}]) +([{_CONTINUOUS_SCRIPT_RANGES}])"
)

# Structural lines whose newlines must survive the merge pass — merging
# a heading into prose destroys the PDF's visible structure.
_RE_HEADING = re.compile(r"^#{1,6}\s")
_RE_LIST = re.compile(r"^(?:[-*•◦▪■]\s|\d+[.)]\s|>\s?)")
_RE_SEPARATOR = re.compile(r"^[\s\-⎯—━═\*\._=·‧…]+$")
_RE_HAS_SEP_CHAR = re.compile(r"[\-⎯—━═_=]")
# Strong sentence terminators — if the previous line ended with one, the
# following line is almost certainly a new sentence / paragraph.
_SENTENCE_ENDS = frozenset('.!?:;—…"」』”’')


def _is_structural_line(line: str) -> bool:
    """Heading / list / separator line that must NOT be merged into prose."""
    line = line.strip()
    if not line:
        return False
    if _RE_HEADING.match(line):
        return True
    if _RE_LIST.match(line):
        return True
    # A line of only decorative chars (⎯, —, -, ═, etc.) is a separator.
    return bool(_RE_SEPARATOR.match(line) and _RE_HAS_SEP_CHAR.search(line))


def _ends_sentence(text: str) -> bool:
    """Did the last non-whitespace char close a sentence?"""
    stripped = text.rstrip()
    if not stripped:
        return False
    return stripped[-1] in _SENTENCE_ENDS


def _rejoin_pdf_lines(text: str) -> str:
    """Fix narrow per-glyph newline artefacts while preserving PDF
    layout. Paired with ``extraction_mode="layout"`` which already
    produces text whose line breaks reflect the source document.

    Only orphan short lines that sit between two longer lines are
    touched — those are almost always positional-jump artefacts pypdf
    split off, e.g.::

        ในห้วง
        อ          ← per-glyph artefact
        วกาศ

    becomes ``ในห้วงอวกาศ``.

    Real line breaks, paragraph breaks, headings, separators, list
    items, and intra-line spaces are all preserved — layout-mode text
    uses them meaningfully. We intentionally do NOT collapse existing
    spaces between continuous-script characters here, because layout
    mode's spacing is the user's original content. Legacy repair for
    pre-layout-mode data lives in ``reflow_legacy_text`` — called by the
    dashboard's Reflow button, never automatically.
    """
    if not text:
        return text

    lines = text.split("\n")
    result: list[str] = []
    i = 0
    while i < len(lines):
        current = lines[i]
        stripped = current.strip()
        if not stripped:
            result.append("")
            i += 1
            continue
        if _is_short_orphan_line(stripped, result, lines, i):
            prev = result[-1]
            last_char = prev[-1] if prev else ""
            first_char = stripped[0]
            if (
                last_char
                and _RE_CONTINUOUS_CHAR.match(last_char)
                and _RE_CONTINUOUS_CHAR.match(first_char)
            ):
                sep = ""
            else:
                sep = " " if prev and not prev.endswith((" ", "\t")) else ""
            result[-1] = prev + sep + stripped
        else:
            result.append(current)
        i += 1
    text = "\n".join(result)
    # Only collapse runaway blank lines; leave horizontal spacing alone
    # because layout mode uses it for visual alignment.
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def reflow_legacy_text(text: str) -> str:
    """Aggressive reflow for PDFs extracted BEFORE layout mode was
    enabled — called only by the dashboard's Reflow button.

    This is the pre-layout-mode rejoin algorithm. It merges short prose
    lines and de-spaces continuous-script runs ("คำ คำ คำ" → "คำคำคำ")
    so legacy data can be cleaned up without re-uploading the source
    PDF. It's aggressive on purpose — legacy data is already mangled,
    so the heuristic doesn't need to be gentle.

    Not used in the ingest path because layout mode already preserves
    structure correctly; running this on fresh data would destroy the
    intentional spacing layout mode produces.
    """
    if not text:
        return text

    # Pass 1 — line-by-line merge with structure awareness.
    lines = text.split("\n")
    result: list[str] = []
    for raw in lines:
        stripped = raw.strip()
        if not stripped:
            if result and result[-1]:
                result.append("")
            continue

        prev = result[-1] if result else ""
        starts_new = (
            not prev
            or _is_structural_line(stripped)
            or _is_structural_line(prev)
            or _ends_sentence(prev)
        )

        if starts_new:
            result.append(stripped)
        else:
            last_char = prev[-1]
            first_char = stripped[0]
            if _RE_CONTINUOUS_CHAR.match(last_char) and _RE_CONTINUOUS_CHAR.match(first_char):
                sep = ""
            else:
                sep = " "
            result[-1] = prev + sep + stripped

    text = "\n".join(result)

    # Pass 2 — re-break inline structural markers.
    text = re.sub(r"(?<!\n) +(#{1,6}\s+)", r"\n\n\1", text)
    text = re.sub(r"[ \t]*([\-⎯—━═_=]{4,})[ \t]*", r"\n\1\n", text)

    # Pass 3 — de-space continuous-script runs.
    prev_txt: str | None = None
    while prev_txt != text:
        prev_txt = text
        text = _RE_SPACE_BETWEEN_CONTINUOUS.sub(r"\1\2", text)

    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _is_short_orphan_line(
    stripped_current: str,
    result_so_far: list[str],
    all_lines: list[str],
    idx: int,
) -> bool:
    """Heuristic: does this short line look like a per-glyph artefact
    that should be merged back into the previous line?

    Criteria:
      - stripped content is short (≤4 chars Thai/CJK, ≤3 chars Latin)
      - previous line exists and has content (not a paragraph break)
      - previous line is longer than this one (prose continuing, not a
        short structural line like a single-char heading)
      - next line (if any) is also longer — indicates the sequence is
        "prose → glyph-orphan → prose" rather than a list or table cell

    Structural lines (headings, list items, separators) are NEVER merged
    even if short, because their newlines matter.
    """
    if _is_structural_line(stripped_current):
        return False
    # Script-based length threshold — Thai glyphs are more information-
    # dense per character than Latin, so 4 Thai chars is a reasonable
    # "orphan" ceiling.
    first_char = stripped_current[0]
    max_len = 4 if _RE_CONTINUOUS_CHAR.match(first_char) else 3
    if len(stripped_current) > max_len:
        return False
    if not result_so_far or not result_so_far[-1].strip():
        return False
    prev_stripped = result_so_far[-1].strip()
    if _is_structural_line(prev_stripped):
        return False
    if len(prev_stripped) <= max_len:
        return False
    # If this is the last line, merging is still fine — it's a trailing
    # glyph orphan.
    if idx + 1 >= len(all_lines):
        return True
    next_stripped = all_lines[idx + 1].strip()
    if not next_stripped:
        return True  # trailing orphan before blank
    if _is_structural_line(next_stripped):
        return False
    return len(next_stripped) > max_len


# ----------------------------------------------------------------------------
# Async extraction + persistence helper (called from request handlers)
# ----------------------------------------------------------------------------

# Total extracted-text storage cap across ALL document memories. 20 MB is
# comfortable for ~100-200 PDFs worth of prose; older memories get LRU-evicted
# when this is exceeded so the DB doesn't grow without bound.
MAX_TOTAL_CHARS = 20_000_000
# Soft cap on row count — LRU eviction triggers here too so very-many-tiny
# files don't blow past ``MAX_TOTAL_CHARS`` before it notices.
MAX_ROWS = 200

# Module-level lock guarding the cap-check / evict / save sequence. Each of
# the underlying DB calls opens its own write connection, so the multi-await
# sequence is otherwise vulnerable to TOCTOU between concurrent uploaders
# (two large docs both pass the size check, both save, both blow past the cap).
# Lazily created so test code that imports the module without an event loop
# doesn't blow up on import.
_persist_lock: Any = None
_extract_sem: Any = None


def _get_persist_lock() -> Any:
    """Return the module-level asyncio.Lock, creating it on first use."""
    global _persist_lock
    if _persist_lock is None:
        import asyncio as _asyncio

        _persist_lock = _asyncio.Lock()
    return _persist_lock


def _get_extract_sem() -> Any:
    """Return the module-level extraction semaphore (cap=2).

    Per-call instantiation (the previous shape) gave each
    ``extract_and_persist`` call its own semaphore, so two simultaneous
    requests each got a fresh cap of 2 → effective concurrency of 4
    across the process, defeating the memory-DoS guard. A module-level
    semaphore caps it globally.
    """
    global _extract_sem
    if _extract_sem is None:
        import asyncio as _asyncio

        _extract_sem = _asyncio.Semaphore(2)
    return _extract_sem


async def extract_and_persist(
    documents: list[Any],
    *,
    db: Any,
    source_conversation_id: str | None = None,
) -> list[dict[str, Any]]:
    """Extract text from each document payload + save to DB.

    Extraction is CPU-bound (``pypdf.extract_text`` can take seconds on
    image-heavy PDFs), so each extraction is dispatched to a worker thread
    via ``asyncio.to_thread`` to keep the event loop responsive. When a user
    attaches multiple files, extractions run **in parallel** via
    ``asyncio.gather`` — the extraction phase is independent per file, so
    five PDFs finish in roughly the time of the slowest single one rather
    than the sum.

    DB writes stay sequential after extraction because the cap-and-evict
    loop reads running totals that change as we go; running them in parallel
    would race on the eviction count.

    Returns a list of ``{"id", "filename", "char_count"}`` dicts for the
    frontend toast. Empty list if nothing was persisted (extraction failed,
    DB unavailable, etc.) — never raises.
    """
    if not documents or db is None:
        return []

    import asyncio

    saved: list[dict[str, Any]] = []

    # Parallel extraction — CPU-bound work in separate threads, then await
    # the gather so we have all results before touching the DB.
    # A module-level semaphore caps concurrent extractions GLOBALLY so a
    # user uploading 5+ huge PDFs can't simultaneously spawn 5 layout-mode
    # workers (each can use GBs of RAM in pypdf), which is a memory-DoS
    # vector. The cap is deliberately conservative — extraction is rare
    # and bursty. (Was previously per-call, defeating the global cap.)
    extract_sem = _get_extract_sem()

    async def _bounded_extract(payload: Any) -> Any:
        async with extract_sem:
            return await asyncio.to_thread(extract_from_payload, payload)

    extractions = await asyncio.gather(
        *(_bounded_extract(payload) for payload in documents),
        return_exceptions=True,
    )

    persist_lock = _get_persist_lock()
    for idx, extracted in enumerate(extractions):
        if isinstance(extracted, BaseException):
            logger.warning(
                "Extraction failed for document %d: %s",
                idx,
                extracted,
            )
            continue
        if extracted is None or not extracted.text.strip():
            continue

        # Reject docs larger than the entire aggregate cap up-front. Without
        # this, the eviction loop below would dutifully delete every existing
        # document memory before discovering the incoming one still doesn't
        # fit — destroying every other user's saved doc to make room for a
        # file that can never fit anyway.
        if extracted.char_count > MAX_TOTAL_CHARS:
            logger.warning(
                "Single document %s exceeds total cap (%d > %d); rejecting",
                extracted.filename,
                extracted.char_count,
                MAX_TOTAL_CHARS,
            )
            continue

        # Serialize the cap-check / evict / save sequence so two concurrent
        # uploads can't both pass the size check and then both write past
        # the cap.
        async with persist_lock:
            # Enforce aggregate caps with LRU eviction. Loop rather than a
            # single check-and-delete because a freshly added big doc could
            # require evicting several older entries. The hard iteration
            # cap protects against a buggy
            # ``delete_oldest_document_memory`` that returns truthy without
            # actually shrinking the table — without the cap this loop
            # would run forever, holding ``persist_lock`` and stalling
            # every concurrent uploader.
            fits = False
            max_eviction_iterations = MAX_ROWS + 16
            try:
                for _ in range(max_eviction_iterations):
                    total = await db.total_document_memories_size()
                    count = await db.count_document_memories()
                    if total + extracted.char_count <= MAX_TOTAL_CHARS and count < MAX_ROWS:
                        fits = True
                        break
                    evicted = await db.delete_oldest_document_memory()
                    if not evicted:
                        # Empty table but still can't fit — shouldn't happen
                        # because the up-front guard already filters this,
                        # but keep the safety bail-out.
                        logger.warning(
                            "Dropping document %s: %d chars exceeds total cap %d",
                            extracted.filename,
                            extracted.char_count,
                            MAX_TOTAL_CHARS,
                        )
                        break
                else:
                    logger.error(
                        "Eviction loop exhausted (%d iterations) for document %s — "
                        "delete_oldest_document_memory may be stuck; bailing.",
                        max_eviction_iterations,
                        extracted.filename,
                    )
            except Exception as e:
                logger.warning("Document memory cap check failed: %s", e)
                continue
            if not fits:
                continue

            try:
                memory_id = await db.save_document_memory(
                    filename=extracted.filename,
                    file_kind=extracted.kind,
                    extracted_text=extracted.text,
                    char_count=extracted.char_count,
                    page_count=extracted.page_count,
                    source_conversation_id=source_conversation_id,
                )
            except Exception as e:
                logger.warning("Failed to save document memory for %s: %s", extracted.filename, e)
                continue

        saved.append(
            {
                "id": memory_id,
                "filename": extracted.filename,
                "file_kind": extracted.kind,
                "char_count": extracted.char_count,
                "page_count": extracted.page_count,
            }
        )
        logger.info(
            "📎 Saved document memory: %s (%d chars, kind=%s)",
            extracted.filename,
            extracted.char_count,
            extracted.kind,
        )

    # If anything was actually persisted, drop the user_context cache so the
    # next AI turn picks up the new docs. Lazy import keeps document_extractor
    # free of the dashboard_common dependency at module-load time (it gets
    # used in test fixtures that don't bring up the chat module).
    if saved:
        try:
            from .dashboard_common import invalidate_user_context_cache

            invalidate_user_context_cache(source_conversation_id)
        except ImportError:
            # Test fixtures may not bring up dashboard_common; in real
            # runtime this is unreachable. Log so a packaging regression
            # (e.g. a renamed module) doesn't hide as "cache never
            # invalidates → AI never sees newly-uploaded docs".
            logger.exception(
                "Failed to invalidate user_context cache after saving "
                "documents; new uploads may not be visible to the next "
                "AI turn until the bot restarts."
            )

    return saved
