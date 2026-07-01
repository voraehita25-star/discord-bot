"""Tests for cogs.ai_core.api.document_extractor.

Focus on the pure-logic helpers (extension detection, base64 decoding,
text normalisation, line rejoining, structural-line classification) plus
the `extract_from_payload` dispatch and `_extract_text` happy path. PDF
and DOCX extraction are smoke-tested with mocked parsers — full
parser fidelity is verified by upstream pypdf / python-docx test
suites and would just retest those libraries here.
"""

from __future__ import annotations

import base64
from unittest.mock import MagicMock, patch

import pytest

from cogs.ai_core.api import document_extractor as de


class TestExtensionOf:
    def test_lowercase_extension(self):
        assert de._extension_of("file.PDF") == ".pdf"

    def test_no_extension(self):
        assert de._extension_of("README") == ""

    def test_path_separator_safe(self):
        assert de._extension_of("a/b/c.docx") == ".docx"

    def test_double_dot(self):
        assert de._extension_of("archive.tar.gz") == ".gz"

    def test_trailing_dot(self):
        assert de._extension_of("file.") == ""


class TestDecodeDataUrl:
    def test_decodes_valid_base64(self):
        raw = b"hello world"
        url = "data:text/plain;base64," + base64.b64encode(raw).decode()
        assert de._decode_data_url(url) == raw

    def test_returns_none_on_missing_comma(self):
        assert de._decode_data_url("data:text/plain;base64") is None

    def test_returns_none_on_wrong_prefix(self):
        assert de._decode_data_url("notdata:text/plain;base64,aGk=") is None

    def test_returns_none_on_corrupt_base64(self):
        assert de._decode_data_url("data:text/plain;base64,!!notbase64!!") is None


class TestNormalise:
    def test_strips_c0_controls(self):
        out = de._normalise("hello\x00world\x07")
        assert out == "helloworld"

    def test_keeps_tabs_and_newlines(self):
        assert de._normalise("a\tb\nc") == "a\tb\nc"

    def test_strips_del_char(self):
        assert de._normalise("hi\x7fbye") == "hibye"

    def test_collapses_blank_lines(self):
        out = de._normalise("a\n\n\n\n\nb")
        assert out == "a\n\nb"

    def test_strips_outer_whitespace(self):
        assert de._normalise("  \n\nhello\n\n  ") == "hello"

    def test_converts_crlf_to_lf(self):
        # The control-char strip keeps \r, so CRLF uploads (the norm on Windows)
        # would otherwise leave a stray \r on every line. _normalise must convert
        # CRLF -> LF up front so no carriage returns survive.
        out = de._normalise("line1\r\nline2\r\nline3")
        assert out == "line1\nline2\nline3"
        assert "\r" not in out

    def test_converts_lone_cr_to_lf(self):
        out = de._normalise("a\rb\rc")
        assert out == "a\nb\nc"
        assert "\r" not in out

    def test_crlf_blank_lines_collapse(self):
        # After CRLF->LF conversion, the \n{3,} blank-line collapse works on
        # CRLF text too (it could not before, since \r\n\r\n\r\n has no run of
        # 3+ consecutive \n).
        out = de._normalise("a\r\n\r\n\r\n\r\nb")
        assert out == "a\n\nb"


class TestEndsSentence:
    @pytest.mark.parametrize("text", ["Hello.", "What?", "Stop!", "Like this:", "End—"])
    def test_recognises_terminators(self, text):
        assert de._ends_sentence(text) is True

    def test_recognises_with_trailing_whitespace(self):
        assert de._ends_sentence("Done.   ") is True

    def test_no_terminator(self):
        assert de._ends_sentence("Just words") is False

    def test_empty_string(self):
        assert de._ends_sentence("") is False

    def test_only_whitespace(self):
        assert de._ends_sentence("   \t\n") is False


class TestIsStructuralLine:
    def test_heading(self):
        assert de._is_structural_line("# Title") is True

    def test_h6(self):
        assert de._is_structural_line("###### deep heading") is True

    def test_bullet(self):
        assert de._is_structural_line("- item") is True
        assert de._is_structural_line("* item") is True
        assert de._is_structural_line("• item") is True

    def test_numbered(self):
        assert de._is_structural_line("1. first") is True
        assert de._is_structural_line("2) second") is True

    def test_quote(self):
        assert de._is_structural_line("> quoted") is True

    def test_separator(self):
        assert de._is_structural_line("---") is True
        assert de._is_structural_line("═══") is True

    def test_prose_not_structural(self):
        assert de._is_structural_line("Just a sentence.") is False

    def test_empty_not_structural(self):
        assert de._is_structural_line("") is False
        assert de._is_structural_line("    ") is False


class TestExtractFromPayloadDispatch:
    def test_returns_none_for_non_dict(self):
        assert de.extract_from_payload("not a dict") is None  # type: ignore[arg-type]
        assert de.extract_from_payload(None) is None  # type: ignore[arg-type]

    def test_returns_none_for_missing_data(self):
        assert de.extract_from_payload({"name": "x.txt"}) is None

    def test_returns_none_for_empty_data(self):
        assert de.extract_from_payload({"name": "x.txt", "data": ""}) is None

    def test_returns_none_for_non_string_data(self):
        assert de.extract_from_payload({"name": "x.txt", "data": 123}) is None

    def test_text_extraction_routes_to_text(self):
        payload = {
            "name": "notes.md",
            "kind": "text",
            "data": "# heading\n\nbody text",
        }
        result = de.extract_from_payload(payload)
        assert result is not None
        assert result.kind == "text"
        assert "heading" in result.text

    def test_pdf_extension_routes_to_pdf(self):
        with patch.object(de, "_extract_pdf") as mock_pdf:
            mock_pdf.return_value = de.ExtractedDocument(
                filename="x.pdf", kind="pdf", text="ok", char_count=2, page_count=1
            )
            result = de.extract_from_payload(
                {"name": "x.pdf", "data": "data:application/pdf;base64,abc"}
            )
            assert result is not None
            mock_pdf.assert_called_once()

    def test_docx_extension_routes_to_docx(self):
        with patch.object(de, "_extract_docx") as mock_docx:
            mock_docx.return_value = de.ExtractedDocument(
                filename="x.docx", kind="docx", text="hi", char_count=2
            )
            result = de.extract_from_payload(
                {
                    "name": "x.docx",
                    "data": "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,abc",
                }
            )
            assert result is not None
            mock_docx.assert_called_once()

    def test_swallow_extractor_exceptions(self):
        with patch.object(de, "_extract_pdf", side_effect=RuntimeError("boom")):
            result = de.extract_from_payload(
                {"name": "x.pdf", "data": "data:application/pdf;base64,abc"}
            )
            assert result is None

    def test_truncates_long_filename(self):
        payload = {"name": "a" * 400 + ".txt", "kind": "text", "data": "ok"}
        result = de.extract_from_payload(payload)
        assert result is not None
        assert len(result.filename) <= 200
        # The extension must survive truncation: it's derived from the full
        # name and re-appended, so a long-but-validly-extensioned file still
        # passes the allowlist guard and routes to text extraction. (Regression
        # guard: slicing the name to [:200] BEFORE deriving the extension drops
        # the suffix and wrongly rejects this attachment -> result is None.)
        assert result.filename.endswith(".txt")
        assert result.kind == "text"

    def test_rejects_dotenv_attachment(self):
        # The temp-file path drops .env (secret-exfil guard); the persistence
        # path must mirror that allowlist so the documented ".env excluded"
        # guarantee holds on BOTH paths. A .env named attachment with kind
        # 'text' must NOT be extracted/persisted.
        payload = {"name": ".env", "kind": "text", "data": "SECRET=hunter2"}
        assert de.extract_from_payload(payload) is None

    def test_rejects_disallowed_extension_even_with_text_kind(self):
        # .exe is outside the allowlist; kind:'text' must not smuggle it past
        # the extension allowlist the temp-file path enforces.
        payload = {"name": "evil.exe", "kind": "text", "data": "payload"}
        assert de.extract_from_payload(payload) is None

    def test_rejects_no_extension_attachment(self):
        # A name with no extension (empty derived ext) is outside the
        # allowlist — the temp-file path drops it, so this path must too.
        payload = {"name": "README", "kind": "text", "data": "hello"}
        assert de.extract_from_payload(payload) is None

    def test_allows_allowlisted_text_extension(self):
        # Sanity: an allowlisted extension still extracts (the allowlist is a
        # rejection filter, not a blanket block).
        payload = {"name": "config.yaml", "kind": "text", "data": "key: value"}
        result = de.extract_from_payload(payload)
        assert result is not None
        assert result.kind == "text"

    def test_allowlist_prefers_canonical_cli_sets(self):
        # The allowlist must be sourced from the canonical sets in
        # dashboard_chat_claude_cli (single source of truth) so the two paths
        # can't drift. Verify the resolver returns those sets when importable.
        from cogs.ai_core.api import dashboard_chat_claude_cli as cli

        de._RESOLVED_DOC_EXT = None  # force re-resolution
        try:
            resolved = de._allowed_doc_extensions()
        finally:
            de._RESOLVED_DOC_EXT = None
        expected = frozenset(cli._SUPPORTED_DOC_BINARY_EXT | cli._SUPPORTED_DOC_TEXT_EXT)
        assert resolved == expected
        assert ".env" not in resolved


class TestExtractText:
    def test_plain_text_payload(self):
        result = de._extract_text("notes.md", "# Title\n\nSome body content.")
        assert result is not None
        assert result.kind == "text"
        assert "Title" in result.text
        assert result.char_count == len(result.text)

    def test_data_url_kept_as_raw_text(self):
        # _extract_text doesn't decode data URLs — text payloads arrive pre-decoded
        # as utf-8 strings. The raw data: prefix becomes part of the body. This
        # pins the contract so anyone refactoring the dispatcher remembers which
        # layer is responsible for decoding.
        body = "# Title\n\nBody"
        url = "data:text/plain;base64," + base64.b64encode(body.encode()).decode()
        result = de._extract_text("notes.md", url)
        assert result is not None
        assert result.text.startswith("data:")

    def test_returns_none_on_empty_payload(self):
        # All-whitespace stripped to empty -> None
        assert de._extract_text("notes.md", "   \n\n\n   ") is None

    def test_caps_at_max_chars(self):
        long_text = "x" * (de.MAX_EXTRACTED_CHARS + 1000)
        result = de._extract_text("big.txt", long_text)
        assert result is not None
        # Truncation appends a "[... truncated]" marker, so allow a small
        # margin above the hard cap rather than insist on strict equality.
        assert result.char_count < de.MAX_EXTRACTED_CHARS + 100
        assert "truncated" in result.text

    def test_pre_sliced_truncation_respects_cap(self):
        # Regression: the pre_sliced branch (input > MAX*4 that _normalise then
        # shrinks back under the cap) appended the 17-char marker WITHOUT a budget
        # slice, so a normalised length just under MAX overflowed the
        # ``len(text) <= MAX_EXTRACTED_CHARS`` invariant by up to 17 chars.
        # Construct exactly that: ~MAX non-whitespace chars followed by a huge
        # collapsible whitespace tail pushing the raw input over MAX*4. After the
        # 2M pre-slice + _normalise (which .strip()s the trailing run) the body is
        # ~MAX-5 chars (in the danger window > MAX-17), so the marker is appended.
        body_len = de.MAX_EXTRACTED_CHARS - 5
        whitespace_tail = " " * (de.MAX_EXTRACTED_CHARS * 4 + 10 - body_len)
        payload = ("x" * body_len) + whitespace_tail
        assert len(payload) > de.MAX_EXTRACTED_CHARS * 4  # ensures pre_sliced=True
        result = de._extract_text("huge.txt", payload)
        assert result is not None
        # The invariant must hold strictly (old code produced MAX + 17 here).
        assert result.char_count <= de.MAX_EXTRACTED_CHARS
        assert result.char_count == len(result.text)
        assert "truncated" in result.text


class TestRejoinPdfLines:
    def test_empty_input(self):
        assert de._rejoin_pdf_lines("") == ""

    def test_preserves_paragraphs(self):
        text = "First paragraph here.\n\nSecond paragraph here."
        out = de._rejoin_pdf_lines(text)
        assert "First paragraph" in out
        assert "Second paragraph" in out

    def test_runs_without_error_on_thai(self):
        # The orphan-merge heuristic kicks in for short Thai glyphs between
        # two longer Thai lines. The exact output depends on _is_short_orphan
        # tuning, so just smoke-test that the function doesn't blow up and
        # produces non-empty output.
        text = "ในห้วงเวลานั้น\nอ\nวกาศก็ขยาย"
        out = de._rejoin_pdf_lines(text)
        assert out

    def test_keeps_structural_lines(self):
        text = "Some prose continues here.\n# Heading\nMore prose follows."
        out = de._rejoin_pdf_lines(text)
        assert "# Heading" in out

    def test_collapses_runaway_blank_lines(self):
        text = "first\n\n\n\n\nlast"
        out = de._rejoin_pdf_lines(text)
        # Max 1 blank between (so 2 newlines)
        assert "\n\n\n" not in out


class TestReflowLegacyText:
    def test_empty_input(self):
        assert de.reflow_legacy_text("") == ""

    def test_collapses_thai_word_spaces(self):
        # Pre-layout-mode PDFs sometimes had a space between every word.
        out = de.reflow_legacy_text("คำ คำ คำ")
        # Aggressive reflow drops the spaces between continuous-script chars.
        assert " " not in out or out.count(" ") < 2

    def test_preserves_latin_word_spaces(self):
        out = de.reflow_legacy_text("hello world how are you")
        # Must NOT collapse spaces between Latin words.
        assert "hello world" in out


class TestExtractedDocument:
    def test_dataclass_basic_fields(self):
        d = de.ExtractedDocument(filename="a.txt", kind="text", text="hi", char_count=2)
        assert d.filename == "a.txt"
        assert d.kind == "text"
        assert d.text == "hi"
        assert d.char_count == 2
        assert d.page_count is None

    def test_dataclass_with_page_count(self):
        d = de.ExtractedDocument(
            filename="a.pdf", kind="pdf", text="hi", char_count=2, page_count=3
        )
        assert d.page_count == 3


class TestPersistLockLazy:
    def test_lock_lazily_created(self):
        # Reset module-level state for a clean test.
        de._persist_lock = None
        import asyncio as _asyncio

        async def check():
            lock = de._get_persist_lock()
            return isinstance(lock, _asyncio.Lock)

        assert _asyncio.run(check()) is True

    def test_lock_is_singleton(self):
        async def check():
            a = de._get_persist_lock()
            b = de._get_persist_lock()
            return a is b

        import asyncio as _asyncio

        assert _asyncio.run(check()) is True


@pytest.mark.asyncio
class TestExtractAndPersistShortcircuits:
    async def test_empty_documents_returns_empty(self):
        result = await de.extract_and_persist([], db=MagicMock())
        assert result == []

    async def test_none_db_returns_empty(self):
        result = await de.extract_and_persist([{"name": "x.txt", "data": "hi"}], db=None)
        assert result == []


class TestExtractPdfMocked:
    """Test the PDF code path with a mocked pypdf reader so we don't need a
    real PDF file in the repo."""

    def _make_url(self, payload: bytes) -> str:
        return "data:application/pdf;base64," + base64.b64encode(payload).decode()

    def test_returns_none_on_bad_data_url(self):
        assert de._extract_pdf("x.pdf", "not-a-data-url") is None

    def test_returns_none_when_pypdf_fails_to_parse(self):
        with patch("pypdf.PdfReader", side_effect=Exception("boom")):
            result = de._extract_pdf("x.pdf", self._make_url(b"%PDF junk"))
            assert result is None

    def test_returns_none_for_encrypted_pdf(self):
        mock_reader = MagicMock()
        mock_reader.is_encrypted = True
        mock_reader.pages = []
        with patch("pypdf.PdfReader", return_value=mock_reader):
            result = de._extract_pdf("locked.pdf", self._make_url(b"%PDF stub"))
            assert result is None

    def test_extracts_simple_text(self):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page body text"
        mock_reader = MagicMock()
        mock_reader.is_encrypted = False
        mock_reader.pages = [mock_page]
        with patch("pypdf.PdfReader", return_value=mock_reader):
            result = de._extract_pdf("hello.pdf", self._make_url(b"%PDF stub"))
            assert result is not None
            assert result.kind == "pdf"
            assert "Page body text" in result.text
            assert result.page_count == 1

    def test_returns_none_when_all_pages_empty(self):
        mock_page = MagicMock()
        mock_page.extract_text.return_value = ""
        mock_reader = MagicMock()
        mock_reader.is_encrypted = False
        mock_reader.pages = [mock_page]
        with patch("pypdf.PdfReader", return_value=mock_reader):
            result = de._extract_pdf("empty.pdf", self._make_url(b"%PDF stub"))
            assert result is None

    def test_falls_back_when_layout_mode_raises(self):
        mock_page = MagicMock()

        def extract_with_modes(extraction_mode=None):
            if extraction_mode == "layout":
                raise RuntimeError("layout broken")
            return "plain mode text"

        mock_page.extract_text.side_effect = extract_with_modes
        mock_reader = MagicMock()
        mock_reader.is_encrypted = False
        mock_reader.pages = [mock_page]
        with patch("pypdf.PdfReader", return_value=mock_reader):
            result = de._extract_pdf("x.pdf", self._make_url(b"%PDF stub"))
            assert result is not None
            assert "plain mode text" in result.text


class TestExtractDocxMocked:
    """Test DOCX extraction with the python-docx Document mocked."""

    def _make_url(self, payload: bytes = b"PK stub") -> str:
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        return f"data:{mime};base64," + base64.b64encode(payload).decode()

    def test_returns_none_on_bad_data_url(self):
        assert de._extract_docx("x.docx", "not-a-data-url") is None

    def test_returns_none_when_disabled(self):
        with patch.object(de, "DOCX_DISABLED", True):
            assert de._extract_docx("x.docx", self._make_url()) is None

    def test_returns_none_when_docx_parse_fails(self):
        if de.DOCX_DISABLED:
            pytest.skip("python-docx unavailable in this env")
        with patch("docx.Document", side_effect=Exception("boom")):
            result = de._extract_docx("x.docx", self._make_url())
            assert result is None

    def test_extracts_paragraph_text(self):
        if de.DOCX_DISABLED:
            pytest.skip("python-docx unavailable in this env")
        # Build a minimal real zip so the zip-bomb preflight passes.
        import io
        import zipfile

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("[Content_Types].xml", b"<x/>")
        zip_url = (
            "data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,"
            + base64.b64encode(buf.getvalue()).decode()
        )
        mock_para1 = MagicMock()
        mock_para1.text = "First paragraph"
        mock_para2 = MagicMock()
        mock_para2.text = "Second paragraph"
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_doc.tables = []
        with patch("docx.Document", return_value=mock_doc):
            result = de._extract_docx("notes.docx", zip_url)
            assert result is not None
            assert result.kind == "docx"
            assert "First paragraph" in result.text
            assert "Second paragraph" in result.text

    def test_rejects_when_decoded_bytes_exceed_cap(self):
        if de.DOCX_DISABLED:
            pytest.skip("python-docx unavailable in this env")
        import io
        import zipfile

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("[Content_Types].xml", b"<x/>")
        zip_bytes = buf.getvalue()
        with (
            patch.object(de, "_MAX_DECODED_DOC_BYTES", len(zip_bytes) - 1),
            patch("docx.Document") as mock_document,
        ):
            result = de._extract_docx("big.docx", self._make_url(zip_bytes))
        assert result is None
        mock_document.assert_not_called()

    def test_docx_parser_does_not_expand_xml_entities(self):
        """XXE / billion-laughs guard: the lxml parser python-docx uses must not
        expand XML entities. We rely on python-docx's own ``resolve_entities=
        False`` parser (verified) rather than the deprecated, no-op
        ``defusedxml.lxml`` monkey-patch — this test pins that protection so a
        future python-docx default change can't silently reintroduce XXE."""
        if de.DOCX_DISABLED:
            pytest.skip("python-docx unavailable in this env")
        from docx.oxml.parser import parse_xml

        xml = b'<?xml version="1.0"?><!DOCTYPE r [<!ENTITY foo "ENTITY_WAS_EXPANDED">]><r>&foo;</r>'
        try:
            el = parse_xml(xml)
        except Exception:
            return  # parser rejected the entity-bearing doc outright — also safe
        # With resolve_entities=False the internal entity is NOT substituted into
        # text; an external SYSTEM entity would likewise never be fetched.
        assert "ENTITY_WAS_EXPANDED" not in (el.text or "")


class TestExtractTextEdgeCases:
    def test_returns_none_for_empty_after_normalise(self):
        # All control chars get stripped → empty → None
        assert de._extract_text("ctrl.txt", "\x00\x01\x02") is None

    def test_preserves_tabs(self):
        result = de._extract_text("tabs.tsv", "a\tb\tc")
        assert result is not None
        assert "\t" in result.text

    def test_truncation_marker(self):
        long = "x" * (de.MAX_EXTRACTED_CHARS + 10)
        result = de._extract_text("big.txt", long)
        assert result is not None
        assert "[... truncated]" in result.text
